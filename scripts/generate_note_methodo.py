"""
Génère la Note Méthodologique du POC (.docx), conforme au modèle imposé :
  1. Dataset retenu (1 p.)
  2. Concepts de l'algorithme récent — SetFit (2 p.)
  3. La modélisation — méthodo, métrique, optimisation (2 p.)
  4. Synthèse des résultats + conclusion (2 p.)
  5. Feature importance globale et locale (2 p.)
  6. Limites et améliorations (1 p.)

Les chiffres proviennent des résultats réels (results/) et de l'analyse
d'interprétabilité (results/interpretability/).

Sortie : Desktop/Projet 9/Note_methodologique.docx
Usage : python scripts/generate_note_methodo.py
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
INTERP = ROOT / "results" / "interpretability"
OUT = ROOT / "Note_methodologique.docx"

BLEU = RGBColor(0x0F, 0x4C, 0x81)
VERT = RGBColor(0x1B, 0x7F, 0x4B)
GRIS = RGBColor(0x55, 0x55, 0x55)


# ── Helpers de mise en forme ──────────────────────────────────────────
def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLEU


def h1(doc, text):
    p = doc.add_heading(level=1)
    for r in p.runs:
        r.font.color.rgb = BLEU
    p.add_run("")
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs or p.runs[0].text == "":
        run = p.add_run(text)
    else:
        p.runs[0].text = text
    run.font.color.rgb = BLEU
    run.font.size = Pt(15)


def h2(doc, text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def body(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def bullet(doc, text, bold_prefix=None, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(size)
    r = p.add_run(text); r.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t


# ── Chargement des données réelles ────────────────────────────────────
def load_interp():
    data = {}
    for name in ["global_tfidf", "local_tfidf", "lime_setfit_local", "lime_setfit_global"]:
        p = INTERP / f"{name}.json"
        if p.exists():
            data[name] = json.loads(p.read_text(encoding="utf-8"))
    return data


# ══════════════════════════════════════════════════════════════════════
def build():
    interp = load_interp()
    doc = Document()

    # Marges un peu réduites pour la densité
    for s in doc.sections:
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

    title(doc, "Note méthodologique — Preuve de concept")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Classification automatique de mails de clubs sportifs\n"
                  "Comparaison TF-IDF, CamemBERT, SetFit, Ministral 8B et système hybride")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRIS

    # ── 1. DATASET ────────────────────────────────────────────────────
    h1(doc, "1. Dataset retenu")
    body(doc,
         "Le POC s'appuie sur deux jeux de données complémentaires, l'un pour la "
         "validation scientifique, l'autre pour le cas d'usage métier.")
    h2(doc, "Dataset métier — Mails synthétiques de clubs sportifs")
    body(doc,
         "Dataset original de 1 800 mails en français (225 par catégorie, parfaitement "
         "équilibré), généré via LLM (Claude Haiku) avec des variations contrôlées de "
         "longueur, ton, urgence, persona expéditeur et sport. Il est étiqueté sur "
         "8 catégories métier : inscription, sponsor, arbitrage-officiels, parent, "
         "fédération, logistique-matchday, indemnités-coachs et divers-administratif. "
         "Split stratifié 80/20 (1 440 entraînement / 360 test).")
    body(doc,
         "Ce dataset matérialise un SaaS pour clubs sportifs dont la brique centrale est "
         "la classification et la priorisation automatique des mails entrants. L'enjeu "
         "critique : un nouveau club arrive avec zéro mail étiqueté.")
    h2(doc, "Dataset de validation — 20 Newsgroups")
    body(doc,
         "Sous-ensemble équilibré de 8 classes (sci.med, sci.space, comp.graphics, "
         "comp.os.ms-windows.misc, rec.autos, rec.sport.hockey, talk.politics.guns, "
         "talk.religion.misc) du benchmark public 20 Newsgroups (~7 200 documents après "
         "nettoyage des en-têtes/citations). Il garantit la reproductibilité scientifique "
         "et permet de situer nos résultats par rapport à un benchmark connu.")
    body(doc,
         "Protocole multi-régimes : pour mesurer la courbe d'apprentissage, on évalue "
         "chaque modèle entraînable à 8, 16, 32, 64 exemples par classe, puis en full data "
         "— ce qui reproduit les différents profils de clubs, du nouveau venu au club établi.")

    doc.add_page_break()

    # ── 2. CONCEPTS SETFIT ────────────────────────────────────────────
    h1(doc, "2. Les concepts de l'algorithme récent : SetFit")
    body(doc,
         "SetFit (Tunstall et al., 2022, arXiv:2209.11055) est un framework de "
         "fine-tuning few-shot, prompt-free, développé par Hugging Face, Intel Labs et "
         "UKP Lab. Il répond à une limite des approches few-shot antérieures (PET, PEFT) : "
         "leur dépendance à des prompts manuels fragiles et à des modèles de plusieurs "
         "milliards de paramètres.")
    h2(doc, "Principe de fonctionnement en deux étapes")
    bullet(doc,
           "Fine-tuning contrastif d'un Sentence Transformer. À partir des rares exemples "
           "étiquetés, on génère des paires de phrases : positives (même catégorie) et "
           "négatives (catégories différentes). Le modèle apprend à rapprocher les "
           "embeddings des paires positives et à éloigner ceux des paires négatives "
           "(apprentissage siamois / contrastif).",
           bold_prefix="Étape 1 — ")
    bullet(doc,
           "Entraînement d'une tête de classification légère (régression logistique) sur "
           "les embeddings produits par le transformer fine-tuné.",
           bold_prefix="Étape 2 — ")
    body(doc,
         "L'astuce centrale est la démultiplication du signal d'apprentissage : avec "
         "16 exemples par classe sur 8 classes (128 exemples), on génère plusieurs milliers "
         "de paires contrastives. C'est ce qui permet à SetFit d'apprendre efficacement là "
         "où un fine-tuning classique manque de données.")
    h2(doc, "Pourquoi ce choix : performances démontrées dans la littérature")
    bullet(doc,
           "Avec seulement 8 exemples étiquetés par classe sur Customer Reviews, SetFit est "
           "compétitif avec un fine-tuning de RoBERTa Large entraîné sur 3 000 exemples.")
    bullet(doc,
           "Sur le benchmark RAFT, SetFit (355M paramètres) atteint 71,3 %, surpassant PET "
           "(69,6 %) et GPT-3 175B (62,7 %), et talonnant T-Few 11B (75,8 %) — soit un "
           "modèle 30× plus petit que GPT-3.")
    bullet(doc,
           "Entraînement en ~30 s pour 0,025 $ (vs 11 min / 0,70 $ pour T-Few 3B), sans GPU "
           "obligatoire. Déploiement local, gratuit à l'inférence, sans dépendance externe.")
    body(doc,
         "Backbone retenu : sentence-transformers/paraphrase-multilingual-mpnet-base-v2, "
         "Sentence Transformer multilingue figurant parmi les références du benchmark "
         "MTEB-fr (Ciancone et al., 2024). Deux variantes de backbone sont également "
         "testées pour évaluer l'apport de la modernité architecturale : ModernBERT "
         "(Warner et al., 2024) et mmBERT (Marone et al., 2025, ModernBERT multilingue "
         "entraîné sur 1 800+ langues).")
    body(doc,
         "Pertinence pour le cas d'usage : SetFit permet le déploiement chez un nouveau "
         "club avec seulement ~10-30 mails étiquetés par catégorie — exactement la "
         "contrainte du « club sans historique ».")

    doc.add_page_break()

    # ── 3. MODELISATION ───────────────────────────────────────────────
    h1(doc, "3. La modélisation")
    h2(doc, "Cinq stratégies comparées, organisées en logique produit")
    add_table(doc,
        ["Niveau", "Stratégie", "Profil de club visé"],
        [
            ["1", "TF-IDF + Régression logistique", "Baseline / club établi"],
            ["2", "CamemBERT fine-tuné", "Club établi (>100 mails/cat.)"],
            ["3", "SetFit (mpnet, ModernBERT, mmBERT)", "Club en onboarding (peu de données)"],
            ["4", "Ministral 8B via API (zero/few-shot)", "Club entrant (zéro donnée)"],
            ["5", "Hybride SetFit + fallback Ministral", "Club exigeant qualité/coût"],
        ])
    h2(doc, "Protocole expérimental")
    bullet(doc, "Modèles entraînables (niv. 1-3) : 5 régimes (8, 16, 32, 64, full) × 5 seeds, "
                "sampling stratifié reproductible — soit la mesure de la courbe d'apprentissage "
                "avec sa variance.")
    bullet(doc, "Ministral (niv. 4) : pas d'entraînement, évaluation sur tout le test en "
                "zero-shot et few-shot 8 (3 seeds pour la variance des exemples in-context).")
    bullet(doc, "Hybride (niv. 5) : construit par-dessus les prédictions SetFit et Ministral ; "
                "balayage de 7 seuils de confiance τ (0,3 à 0,9).")
    bullet(doc, "Total : ~380 expériences, réalisées en local sur GPU RTX 5070.")
    h2(doc, "Métrique d'évaluation retenue")
    body(doc,
         "La métrique principale est le F1 macro : moyenne non pondérée des F1 par classe. "
         "Ce choix est adapté à des catégories potentiellement déséquilibrées en production "
         "(une petite catégorie compte autant qu'une grande), ce qui évite qu'un modèle "
         "performant seulement sur les classes fréquentes paraisse bon. L'accuracy, le temps "
         "et le coût d'inférence par mail, ainsi que la taille du modèle, complètent l'analyse.")
    h2(doc, "Démarche d'optimisation")
    bullet(doc, "SetFit : 1 epoch contrastif, plafond de paires (max_steps) pour borner le "
                "coût quadratique sur les gros régimes, body LR 1e-5 / head LR 1e-2, "
                "longueur de séquence 128 tokens (suffisante pour les mots discriminants).")
    bullet(doc, "CamemBERT : 3 epochs, learning rate 2e-5, longueur 256 tokens.")
    bullet(doc, "TF-IDF : n-grammes (1,2), sublinear TF, min_df=2, class_weight équilibré.")
    bullet(doc, "Ministral : température 0 (déterministe), sortie JSON structurée parsée "
                "automatiquement, prompt adapté au domaine de chaque dataset.")
    bullet(doc, "Hybride : le seuil τ est le levier d'optimisation produit — il arbitre "
                "explicitement entre qualité, coût API et latence.")

    doc.add_page_break()

    # ── 4. SYNTHESE RESULTATS ─────────────────────────────────────────
    h1(doc, "4. Synthèse des résultats et conclusion")
    h2(doc, "F1 macro — dataset métier (mails de clubs sportifs)")
    add_table(doc,
        ["Modèle", "8-shot", "Full data"],
        [
            ["TF-IDF + LR", "0,84", "0,99"],
            ["CamemBERT", "0,10", "0,99"],
            ["SetFit (mpnet)", "0,95", "0,99"],
            ["SetFit (mmBERT)", "0,93", "0,99"],
            ["Ministral 8B", "0,79 (zero-shot) / 0,90 (few-shot 8)", "—"],
        ])
    h2(doc, "F1 macro — dataset de validation (20 Newsgroups)")
    add_table(doc,
        ["Modèle", "8-shot", "Full data"],
        [
            ["TF-IDF + LR", "0,38", "0,80"],
            ["CamemBERT", "0,08", "0,79"],
            ["SetFit (mpnet)", "0,78", "0,84"],
            ["SetFit (ModernBERT-EN)", "0,64", "0,80"],
            ["Ministral 8B", "0,82 (zero-shot)", "—"],
        ])
    h2(doc, "Lectures clés")
    bullet(doc, "SetFit domine nettement le few-shot : 0,95 dès 8 exemples/classe sur les "
                "emails, là où la baseline plafonne à 0,84 et CamemBERT s'effondre à 0,10. "
                "C'est la réponse directe au problème du nouveau club.", bold_prefix="Few-shot — ")
    bullet(doc, "CamemBERT est data-hungry : nul en few-shot, il rejoint le peloton (0,99) "
                "seulement en full data. À réserver aux clubs à fort historique.", bold_prefix="Full data — ")
    bullet(doc, "En full data, TF-IDF égale les modèles lourds (0,99) : le vocabulaire métier "
                "est très discriminant. C'est ce qui justifie son déploiement (modèle 500× plus "
                "léger, sans GPU).", bold_prefix="Baseline — ")
    bullet(doc, "Système hybride : à τ=0,8, F1 = 0,997 en n'escaladant que ~1 % des mails "
                "vers le LLM payant. Le seuil τ devient un levier produit qualité/coût.",
                bold_prefix="Hybride — ")
    bullet(doc, "Ministral confirme la sensibilité des LLM à la sémantique des labels : "
                "excellent quand les catégories sont parlantes (emails), il avait initialement "
                "chuté sur Newsgroups à cause d'un bug d'étiquetage dans notre vérité-terrain "
                "— corrigé, il remonte à 0,82. Le LLM a en réalité révélé le bug.",
                bold_prefix="LLM — ")
    h2(doc, "Matrice de décision produit (conclusion)")
    add_table(doc,
        ["Profil de club", "Données", "Stratégie recommandée"],
        [
            ["Club entrant", "0 mail", "Ministral 8B zero-shot"],
            ["Club en onboarding", "~10-30/cat.", "SetFit few-shot"],
            ["Club établi", ">100/cat.", "TF-IDF ou CamemBERT (équivalents, 0,99)"],
            ["Club exigeant", "Variable", "Hybride SetFit + Ministral (τ ajustable)"],
        ])
    body(doc,
         "Conclusion : il n'existe pas un meilleur modèle universel, mais une stratégie "
         "optimale par profil de club. La technique récente (SetFit) apporte un gain décisif "
         "précisément là où les approches classiques échouent — le démarrage à faibles "
         "données — tout en restant locale et économique.")

    doc.add_page_break()

    # ── 5. FEATURE IMPORTANCE ─────────────────────────────────────────
    h1(doc, "5. Feature importance globale et locale")
    body(doc,
         "L'interprétabilité est analysée sur les deux familles de modèles : TF-IDF + LR "
         "(nativement interprétable, ses coefficients sont directement des poids de mots) et "
         "SetFit (boîte noire, expliqué via LIME — approximation locale par un modèle "
         "linéaire autour de chaque prédiction).")

    h2(doc, "Importance globale — TF-IDF + LR")
    body(doc,
         "Pour chaque catégorie, les mots de plus fort coefficient révèlent un vocabulaire "
         "métier parfaitement cohérent, ce qui explique la robustesse du modèle :")
    if "global_tfidf" in interp:
        g = interp["global_tfidf"]
        rows = []
        for cls in ["inscription", "sponsor", "arbitrage-officiels", "parent",
                    "federation", "logistique-matchday", "indemnites-coachs",
                    "divers-administratif"]:
            if cls in g:
                mots = ", ".join(x["mot"] for x in g[cls][:5])
                rows.append([cls, mots])
        add_table(doc, ["Catégorie", "Mots les plus discriminants"], rows)
    if (INTERP / "global_tfidf.png").exists():
        doc.add_paragraph()
        doc.add_picture(str(INTERP / "global_tfidf.png"), width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figure — Mots les plus discriminants par catégorie (coefficients TF-IDF + LR).")
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRIS

    h2(doc, "Importance globale — SetFit (via LIME)")
    if "lime_setfit_global" in interp:
        mots = ", ".join(x["mot"] for x in interp["lime_setfit_global"]["mots_influents_globaux"][:12])
        body(doc,
             f"L'agrégation des poids LIME sur un échantillon de mails fait ressortir les "
             f"mots les plus influents sur les décisions de SetFit : {mots}. On retrouve un "
             f"vocabulaire métier discriminant (attestation, arbitres, fédération, indemnités…), "
             f"ce qui confirme que SetFit s'appuie sur des indices sémantiquement pertinents.")

    h2(doc, "Importance locale — explication d'une prédiction")
    if "local_tfidf" in interp and "lime_setfit_local" in interp:
        lt = interp["local_tfidf"]
        ls = interp["lime_setfit_local"]
        body(doc,
             f"Sur un même mail (prédit « {lt['prediction']} »), les deux modèles "
             f"convergent vers les mêmes indices :")
        tfidf_mots = ", ".join(x["mot"] for x in lt["contributions"][:5])
        lime_mots = ", ".join(x["mot"] for x in ls["mots_influents"][:5] if x["poids"] > 0)
        bullet(doc, f"TF-IDF + LR : {tfidf_mots}.", bold_prefix="Contributions ")
        bullet(doc, f"SetFit (LIME) : {lime_mots}.", bold_prefix="Mots influents ")
        body(doc,
             "Cette convergence entre un modèle linéaire transparent et une boîte noire "
             "explicitée par LIME renforce la confiance dans les deux approches : elles "
             "fondent leurs décisions sur les mêmes signaux métier (« vestiaires », "
             "« terrain », « match », « samedi » pour la logistique d'un match).")

    doc.add_page_break()

    # ── 6. LIMITES ────────────────────────────────────────────────────
    h1(doc, "6. Limites et améliorations possibles")
    h2(doc, "Limites de l'approche")
    bullet(doc, "Dataset synthétique : les mails sont générés par LLM. Leur vocabulaire est "
                "peut-être plus « propre » et plus séparable que des mails réels, ce qui peut "
                "surévaluer les performances (notamment le 0,99 de TF-IDF en full data).")
    bullet(doc, "SetFit, modèle déployé final : c'est TF-IDF qui est déployé (léger, cloud), "
                "non SetFit. Le champion few-shot n'est donc pas servi en production faute "
                "d'infrastructure GPU — un écart entre l'optimum scientifique et l'optimum "
                "d'ingénierie, assumé explicitement.")
    bullet(doc, "Confiance des LLM peu fiable : la confiance de Ministral est auto-déclarée "
                "(souvent ~0,5), contrairement aux probabilités calibrées de TF-IDF/SetFit. "
                "C'est pourquoi l'hybride se fonde sur la confiance de SetFit, pas du LLM.")
    bullet(doc, "Coût quadratique de SetFit : le nombre de paires explose avec les données ; "
                "on a dû plafonner les steps, ce qui peut légèrement brider le full data.")
    bullet(doc, "LIME est une approximation locale et stochastique : les explications varient "
                "d'une exécution à l'autre et ne garantissent pas la fidélité parfaite au modèle.")
    h2(doc, "Améliorations envisageables")
    bullet(doc, "Valider sur des mails réels (anonymisés) pour confirmer la généralisation et "
                "recalibrer les attentes de performance.")
    bullet(doc, "Déployer SetFit via un service dédié (Hugging Face Inference Endpoint ou "
                "Space) appelé par le dashboard, pour servir le meilleur modèle few-shot.")
    bullet(doc, "Calibration des probabilités (Expected Calibration Error, reliability "
                "diagrams) pour fiabiliser le score de confiance, critique pour la priorisation "
                "et le routage en validation humaine.")
    bullet(doc, "Interprétabilité : compléter LIME par SHAP (valeurs de Shapley, plus "
                "rigoureuses) et analyser systématiquement les cas de désaccord entre modèles.")
    bullet(doc, "Enrichir le few-shot par de l'augmentation de données ou un active learning "
                "(faire étiqueter en priorité les mails les plus incertains).")

    doc.save(OUT)
    print(f"Note méthodologique générée : {OUT}")


if __name__ == "__main__":
    build()
