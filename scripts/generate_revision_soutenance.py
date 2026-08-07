"""
Génère la note de révision pour la soutenance.

Trois parties :
  A. Glossaire — chaque terme technique du projet expliqué simplement
  B. Slide par slide — le message à faire passer, et les questions probables
  C. Les questions transversales redoutables + l'antisèche chiffrée

Sortie : Desktop/Projet 9/Revision_soutenance.docx
Usage  : python scripts/generate_revision_soutenance.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).parents[1] / "Revision_soutenance.docx"

BLEU = RGBColor(0x0F, 0x4C, 0x81)
VERT = RGBColor(0x1B, 0x7F, 0x4B)
ORANGE = RGBColor(0xB8, 0x5C, 0x00)
GRIS = RGBColor(0x55, 0x55, 0x55)
NOIR = RGBColor(0x1A, 0x1A, 0x1A)


# ── Briques de mise en forme ─────────────────────────────────────────────────
def titre(doc, texte, taille=22):
    p = doc.add_paragraph()
    r = p.add_run(texte)
    r.bold = True
    r.font.size = Pt(taille)
    r.font.color.rgb = BLEU
    return p


def h1(doc, texte):
    p = doc.add_heading(level=1)
    r = p.add_run(texte)
    r.font.color.rgb = BLEU
    r.font.size = Pt(16)
    return p


def h2(doc, texte):
    p = doc.add_heading(level=2)
    r = p.add_run(texte)
    r.font.color.rgb = NOIR
    r.font.size = Pt(12.5)
    return p


def para(doc, texte, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(texte)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, texte, prefixe=None, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    if prefixe:
        r = p.add_run(prefixe)
        r.bold = True
        r.font.size = Pt(size)
    r2 = p.add_run(texte)
    r2.font.size = Pt(size)
    return p


def champ(doc, label, texte, couleur=VERT, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(f"{label} : ")
    r.bold = True
    r.font.color.rgb = couleur
    r.font.size = Pt(size)
    r2 = p.add_run(texte)
    r2.font.size = Pt(size)
    return p


def qr(doc, question, reponse):
    """Un couple question du jury / réponse à donner."""
    p = doc.add_paragraph()
    r = p.add_run(f"Q. {question}")
    r.bold = True
    r.font.color.rgb = ORANGE
    r.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"R. {reponse}")
    r2.font.size = Pt(10.5)
    p2.paragraph_format.space_after = Pt(10)
    return p2


def terme(doc, mot, definition, exemple=None):
    """Entrée de glossaire : le mot, sa définition, et si utile une illustration."""
    p = doc.add_paragraph()
    r = p.add_run(mot)
    r.bold = True
    r.font.color.rgb = BLEU
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(definition)
    r2.font.size = Pt(10.5)
    p2.paragraph_format.space_after = Pt(2)

    if exemple:
        p3 = doc.add_paragraph()
        r3 = p3.add_run(exemple)
        r3.italic = True
        r3.font.size = Pt(10)
        r3.font.color.rgb = GRIS
        p3.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════
def partie_a_glossaire(doc):
    h1(doc, "PARTIE A — Le glossaire : tous les termes du projet")
    para(doc, "Chaque terme que tu as signalé, expliqué comme tu pourrais l'expliquer "
              "à un jury non spécialiste. Si tu ne devais réviser qu'une partie, c'est celle-ci.",
         italic=True, color=GRIS, size=10)

    # ── A1 ────────────────────────────────────────────────────────────────
    h2(doc, "A1. Les cinq stratégies comparées")

    terme(doc, "TF-IDF + Régression logistique (la baseline)",
          "Deux briques enchaînées. TF-IDF transforme un mail en une liste de nombres : "
          "pour chaque mot du vocabulaire, un score qui monte si le mot est fréquent DANS CE "
          "MAIL (Term Frequency) et qui descend s'il est banal dans TOUS les mails (Inverse "
          "Document Frequency). Les mots comme « le » ou « bonjour » sont donc écrasés, les "
          "mots rares et discriminants comme « vestiaires » ressortent. La régression "
          "logistique apprend ensuite un poids par mot et par catégorie, puis additionne "
          "ces poids pour décider.",
          "Formule à retenir : TF-IDF = « ce mot est-il caractéristique de CE document ? ». "
          "C'est de la statistique de comptage, aucun réseau de neurones. D'où sa légèreté : "
          "2,3 Mo et pas de carte graphique.")

    terme(doc, "CamemBERT (fine-tuning classique)",
          "Un modèle de langue français pré-entraîné (110 millions de paramètres) sur d'énormes "
          "quantités de texte. Le « fine-tuning » consiste à reprendre ce modèle déjà savant et "
          "à le ré-entraîner sur nos 8 catégories. Il comprend le sens des phrases, pas "
          "seulement les mots présents.",
          "Sa faiblesse dans mon POC : il a 110 millions de paramètres à ajuster. Avec "
          "64 mails d'entraînement, il n'a pas de quoi les régler — d'où l'effondrement à 0,10.")

    terme(doc, "SetFit",
          "Une méthode d'entraînement en deux temps, conçue pour les situations où l'on a très "
          "peu d'exemples étiquetés. Étape 1 : on prend un modèle qui sait mesurer la "
          "ressemblance entre phrases, et on l'affine sur des PAIRES de mails (« ces deux-là "
          "sont de la même catégorie » / « ceux-là non »). Étape 2 : on entraîne une simple "
          "régression logistique sur les représentations produites.",
          "L'astuce : 64 mails donnent seulement 64 exemples pour un fine-tuning classique, "
          "mais des milliers de paires pour SetFit. C'est cette démultiplication qui fait toute "
          "la différence.")

    terme(doc, "« Framework de fine-tuning few-shot et prompt-free »",
          "Décomposons. FRAMEWORK : une méthode outillée, avec une librairie prête à l'emploi. "
          "FINE-TUNING : on part d'un modèle pré-entraîné qu'on spécialise, plutôt que de partir "
          "de zéro. FEW-SHOT : on n'a que quelques exemples par catégorie (typiquement 8 à 64). "
          "PROMPT-FREE : contrairement aux méthodes qui exigent de rédiger une phrase à trous "
          "bien tournée, SetFit n'a besoin d'aucune consigne textuelle — juste des exemples "
          "étiquetés.",
          "Pourquoi « prompt-free » est un argument fort : rédiger un bon prompt est un art "
          "fragile, un mot changé peut faire chuter la performance. SetFit supprime cette "
          "dépendance.")

    terme(doc, "Ministral 8B utilisé comme classifieur (LLM-as-classifier)",
          "On n'entraîne rien du tout. On envoie le mail à un modèle génératif via une API, "
          "avec une consigne en français : « Voici 8 catégories, classe ce mail, réponds en "
          "JSON ». Le modèle répond avec une catégorie.",
          "Intérêt produit : opérationnel en cinq minutes, sans une seule donnée étiquetée. "
          "Coût : environ 0,000022 $ par mail, et 400 à 600 ms de latence.")

    terme(doc, "Le système hybride (routage par confiance)",
          "SetFit traite tous les mails localement et gratuitement. Chaque prédiction "
          "s'accompagne d'un score de confiance. Si ce score descend sous un seuil τ (tau), on "
          "considère que SetFit hésite, et on envoie ce mail-là — et seulement celui-là — au "
          "modèle payant.",
          "Le seuil τ est un curseur produit : plus on le monte, plus on escalade, plus on paie, "
          "meilleure est la qualité. À τ=0,8 sur données complètes : F1 0,997 pour 1,1 % "
          "d'escalade seulement.")

    # ── A2 ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h2(doc, "A2. Le vocabulaire du few-shot")

    terme(doc, "Few-shot, zero-shot",
          "FEW-SHOT : apprendre à partir de très peu d'exemples étiquetés. ZERO-SHOT : classer "
          "sans aucun exemple, uniquement à partir d'une description de la tâche. Dans mon POC, "
          "seul Ministral peut faire du zero-shot, puisque c'est le seul qui n'a pas besoin "
          "d'entraînement.")

    terme(doc, "PET (Pattern-Exploiting Training)",
          "Une méthode few-shot antérieure à SetFit. Elle transforme la classification en un jeu "
          "de texte à trous : au lieu de demander « quelle catégorie ? », on présente au modèle "
          "« Ce mail parle de ____ » et on regarde quel mot il propose de mettre dans le blanc. "
          "Il faut donc écrire à la main une phrase-modèle (le « pattern ») et une table de "
          "correspondance mot → catégorie (le « verbalizer »).",
          "Sa faiblesse : tout repose sur la qualité de ces phrases écrites à la main. "
          "C'est exactement ce que SetFit élimine en étant prompt-free.")

    terme(doc, "PEFT (Parameter-Efficient Fine-Tuning)",
          "Ce n'est pas une méthode unique mais une famille : au lieu de ré-entraîner les "
          "millions de paramètres d'un gros modèle, on les gèle et on n'entraîne qu'un tout "
          "petit nombre de paramètres ajoutés. Cela rend le fine-tuning de très gros modèles "
          "abordable. T-Few, cité dans l'article SetFit, appartient à cette famille.",
          "Attention à ne pas confondre : PET est une technique de prompt, PEFT une technique "
          "d'optimisation. Les deux sont citées comme « l'état de l'art d'avant SetFit ».")

    terme(doc, "Le benchmark RAFT",
          "Real-world Annotated Few-shot Tasks. Un banc d'essai public conçu pour évaluer les "
          "méthodes few-shot dans des conditions réalistes : 11 tâches de classification "
          "issues de cas réels, avec seulement 50 exemples d'entraînement par tâche. Il existe "
          "un classement public, et une « performance humaine » de référence.",
          "Le chiffre à citer : SetFit (355 M de paramètres) obtient 71,3 %, contre 62,7 % pour "
          "GPT-3 (175 milliards de paramètres). Un modèle 30 fois plus petit fait mieux.")

    terme(doc, "Hugging Face, Intel Labs, UKP Lab",
          "Les trois organisations qui ont co-signé l'article SetFit. HUGGING FACE : l'entreprise "
          "qui héberge la principale plateforme de partage de modèles d'IA, et qui publie les "
          "librairies transformers et setfit. INTEL LABS : la division recherche d'Intel. "
          "UKP LAB : le laboratoire de traitement du langage de l'université technique de "
          "Darmstadt, en Allemagne — c'est le laboratoire d'origine de Sentence-BERT, la brique "
          "sur laquelle SetFit s'appuie.",
          "À dire si on te demande la crédibilité de la source : c'est une collaboration "
          "industrie + recherche académique, publiée sur arXiv et adoptée dans une librairie "
          "open source largement utilisée.")

    # ── A3 ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h2(doc, "A3. Le protocole expérimental")

    terme(doc, "F1 macro",
          "Le F1 d'une catégorie combine deux choses : la précision (parmi les mails que j'ai "
          "classés « sponsor », combien le sont vraiment ?) et le rappel (parmi les vrais mails "
          "« sponsor », combien ai-je retrouvés ?). Le F1 est leur moyenne harmonique — elle "
          "punit tout déséquilibre entre les deux. Le F1 MACRO est ensuite la moyenne SIMPLE "
          "des F1 des 8 catégories.")

    terme(doc, "Pourquoi « une petite catégorie pèse autant qu'une grande »",
          "Parce que la moyenne est simple, non pondérée par les effectifs. Prenons un exemple "
          "chiffré : 8 catégories, dont 7 comptent 100 mails et une seule en compte 10. Un "
          "modèle qui réussit parfaitement les 7 grosses et rate totalement la petite obtient "
          "une exactitude (accuracy) de 700/710 = 98,6 % — ce qui paraît excellent. Mais son "
          "F1 macro vaut (7 × 1,0 + 0) / 8 = 0,875, ce qui révèle immédiatement qu'une "
          "catégorie entière est ignorée.",
          "L'enjeu métier : chez un vrai club, « indemnités-coachs » est sans doute bien plus "
          "rare que « parent ». Avec l'accuracy, un modèle pourrait ignorer complètement les "
          "indemnités sans que ça se voie. Le F1 macro l'interdit.")

    terme(doc, "Seed (graine aléatoire)",
          "Un nombre qui initialise le générateur de hasard. Beaucoup d'étapes comportent de "
          "l'aléatoire : quels 8 mails je tire pour l'entraînement, comment j'initialise les "
          "poids, dans quel ordre je présente les exemples. En fixant la seed, ces tirages "
          "deviennent reproductibles à l'identique.",
          "J'ai utilisé 5 seeds (42, 123, 456, 789, 1024) par régime. Objectif : mesurer si un "
          "bon score est solide ou s'il tenait à un tirage chanceux. C'est la bande de "
          "dispersion autour de mes courbes.")

    terme(doc, "Échantillonnage stratifié reproductible",
          "STRATIFIÉ : quand je tire 8 exemples par catégorie, je garantis d'en avoir exactement "
          "8 de chaque — pas 15 « parent » et 2 « sponsor » par accident du tirage. On respecte "
          "la structure des classes. REPRODUCTIBLE : le tirage est piloté par la seed, donc "
          "n'importe qui relançant mon code avec la seed 42 obtient exactement les mêmes mails.",
          "Pourquoi c'est important : sans stratification, la comparaison entre modèles serait "
          "faussée — un modèle pourrait sembler meilleur simplement parce qu'il a reçu un "
          "tirage plus équilibré.")

    terme(doc, "Régime",
          "Mon mot pour désigner une taille de jeu d'entraînement : le régime « 8 » signifie "
          "8 exemples étiquetés par catégorie, soit 64 mails au total. Les cinq régimes "
          "(8, 16, 32, 64, données complètes) simulent la maturité croissante d'un club.")

    terme(doc, "Epoch, et « epoch contrastif »",
          "Une EPOCH est un passage complet sur l'ensemble des données d'entraînement. Trois "
          "epochs = le modèle voit trois fois chaque exemple. Pour SetFit, l'unité n'est pas le "
          "mail mais la PAIRE de mails : une « epoch contrastive » est donc un passage complet "
          "sur toutes les paires générées.",
          "C'est pour ça qu'une seule epoch suffit à SetFit : cette unique passe traite déjà "
          "des milliers de paires. Là où CamemBERT, avec 3 epochs sur 64 mails, ne fait que "
          "12 pas d'optimisation.")

    # ── A4 ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h2(doc, "A4. Les hyperparamètres, ligne par ligne")
    para(doc, "C'est la diapositive de ta note méthodologique la plus susceptible de "
              "déclencher une question précise. Voici chaque réglage traduit.",
         italic=True, color=GRIS, size=10)

    para(doc, "SetFit", color=VERT, size=11)
    bullet(doc, "un seul passage sur les paires générées. Suffisant car ces paires se "
                "comptent en milliers ; au-delà, le modèle commence à sur-apprendre.",
           "1 epoch contrastif — ")
    bullet(doc, "le nombre de paires croît comme le carré du nombre d'exemples. Avec toutes "
                "les données, cela deviendrait ingérable. Ce plafond (fixé à 1 500 pas) borne "
                "le temps de calcul et garde les régimes comparables entre eux. La convergence "
                "est atteinte bien avant.",
           "plafond de paires (max_steps) — ")
    bullet(doc, "le taux d'apprentissage du transformer lui-même. Volontairement petit : ce "
                "modèle sait déjà représenter le langage, on veut l'ajuster en douceur, pas lui "
                "faire oublier ce qu'il sait.",
           "body learning rate 1e-5 — ")
    bullet(doc, "le taux d'apprentissage de la régression logistique finale. Mille fois plus "
                "grand, car cette tête part de zéro : elle doit apprendre vite et n'a rien à "
                "oublier.",
           "head learning rate 1e-2 — ")
    bullet(doc, "on ne lit que les 128 premiers mots-morceaux de chaque mail. Les mots "
                "discriminants (« vestiaires », « licence ») apparaissent tôt ; lire plus long "
                "coûterait du calcul sans gain.",
           "longueur de séquence 128 tokens — ")

    para(doc, "CamemBERT", color=VERT, size=11)
    bullet(doc, "trois passages sur les données. Constant sur tous les régimes, pour que la "
                "courbe d'apprentissage reste interprétable — c'est un choix de comparabilité, "
                "et c'est aussi ce qui explique la non-convergence à 8 exemples.",
           "3 epochs — ")
    bullet(doc, "valeur standard pour le fine-tuning d'un BERT. Plus haut, le modèle "
                "pré-entraîné se dégrade ; plus bas, il n'apprend pas la tâche.",
           "learning rate 2e-5 — ")
    bullet(doc, "deux fois plus que SetFit, car ici le modèle entier lit le mail pour décider ; "
                "on lui laisse plus de contexte.",
           "longueur 256 tokens — ")

    para(doc, "TF-IDF", color=VERT, size=11)
    bullet(doc, "on compte les mots seuls ET les paires de mots consécutifs. « club sportif » "
                "devient une caractéristique à part entière, plus informative que « club » et "
                "« sportif » séparément.",
           "n-grammes (1,2) — ")
    bullet(doc, "on applique un logarithme au comptage. Un mot présent 10 fois n'est pas "
                "10 fois plus important qu'un mot présent 1 fois — cela évite qu'une répétition "
                "écrase tout le reste.",
           "sublinear TF — ")
    bullet(doc, "on ignore les mots qui n'apparaissent que dans un seul document. Ce sont "
                "presque toujours des fautes de frappe ou des noms propres, sans valeur "
                "généralisable.",
           "min_df = 2 — ")
    bullet(doc, "le modèle compense automatiquement les catégories moins représentées en leur "
                "donnant plus de poids. Ici mon corpus est équilibré, mais c'est une sécurité "
                "pour un déploiement réel où il ne le sera pas.",
           "class_weight équilibré — ")

    # ── A5 ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h2(doc, "A5. L'interprétabilité")

    terme(doc, "LIME",
          "Local Interpretable Model-agnostic Explanations. Une méthode pour expliquer une "
          "prédiction d'un modèle opaque. Le principe : on prend le mail, on en fabrique des "
          "centaines de versions abîmées en retirant des mots au hasard, on demande au modèle "
          "de classer chacune, puis on regarde quels mots faisaient basculer la décision quand "
          "ils disparaissaient. On ajuste enfin un petit modèle linéaire, lisible, autour de "
          "ce seul mail.",
          "« Local » : l'explication vaut pour CE mail, pas pour le modèle en général. "
          "« Model-agnostic » : ça marche sur n'importe quel modèle, on n'a pas besoin de "
          "regarder à l'intérieur.")

    terme(doc, "Pourquoi LIME est stochastique",
          "Parce que le tirage des mots retirés est aléatoire. Deux exécutions sur le même mail "
          "ne produisent pas exactement la même liste de mots influents — l'ordre peut changer, "
          "un mot peut entrer ou sortir du top 5.",
          "Comment le maîtriser, si on te le demande : fixer la graine aléatoire, et augmenter "
          "le nombre de perturbations générées. C'est le compromis stabilité / temps de calcul. "
          "C'est aussi pourquoi je le mentionne comme une limite dans ma note.")

    terme(doc, "SHAP, et pourquoi ce serait mieux",
          "SHAP repose sur les valeurs de Shapley, un concept issu de la théorie des jeux "
          "coopératifs : on répartit équitablement le « gain » (ici la prédiction) entre les "
          "« joueurs » (ici les mots), en considérant toutes les combinaisons possibles. Cela "
          "donne des garanties mathématiques que LIME n'a pas : cohérence, additivité des "
          "contributions.",
          "Le prix à payer : c'est nettement plus coûteux à calculer. C'est pour cela que je le "
          "cite en amélioration et non en réalisation.")

    # ── A6 ────────────────────────────────────────────────────────────────
    h2(doc, "A6. La confiance et sa calibration")

    terme(doc, "Confiance auto-déclarée : pourquoi elle n'est pas fiable",
          "Quand TF-IDF ou SetFit annoncent 0,87, ce nombre sort d'un calcul : la régression "
          "logistique produit une vraie distribution de probabilité sur les 8 catégories. Quand "
          "Ministral annonce 0,87, il a simplement ÉCRIT ce nombre dans son JSON — c'est un mot "
          "généré comme les autres, une imitation de ce à quoi ressemble une réponse confiante. "
          "Aucun calcul de probabilité derrière.",
          "Dans mes mesures, Ministral répondait très souvent 0,5 quelle que soit la situation. "
          "C'est précisément pour cela que mon système hybride se fonde sur la confiance de "
          "SetFit et jamais sur celle du LLM. À dire tel quel si on t'interroge sur l'hybride.")

    terme(doc, "Calibration",
          "Un modèle est bien calibré si ses scores de confiance disent la vérité : parmi tous "
          "les mails où il annonce 80 % de confiance, il devrait avoir raison environ 80 % du "
          "temps. S'il a raison 95 % du temps, il est trop modeste ; s'il n'a raison que 60 % "
          "du temps, il est trop sûr de lui (le cas le plus fréquent et le plus dangereux).")

    terme(doc, "Ce que veut dire « mesurer la calibration » concrètement",
          "Trois étapes. (1) On regroupe toutes les prédictions par tranche de confiance : "
          "celles entre 0,5 et 0,6, entre 0,6 et 0,7, etc. (2) Pour chaque tranche, on compare "
          "la confiance moyenne annoncée au taux de bonnes réponses réellement observé. "
          "(3) L'ECE (Expected Calibration Error) est la moyenne des écarts entre les deux, "
          "pondérée par le nombre de prédictions dans chaque tranche. Zéro = parfaitement "
          "calibré. Le RELIABILITY DIAGRAM est le graphique correspondant : confiance annoncée "
          "en abscisse, exactitude observée en ordonnée. Un modèle parfait suit la diagonale.",
          "Pourquoi c'est critique dans MON projet : tout le système hybride repose sur le seuil "
          "τ appliqué à la confiance de SetFit. Si cette confiance est mal calibrée, je route "
          "vers le LLM payant les mauvais mails — je paie sans gagner en qualité. C'est aussi "
          "ce qui permettrait de dire à un club « ces 30 mails-là méritent une relecture "
          "humaine ». Je ne l'ai pas mesuré : c'est une limite assumée, pas un oubli.")

    # ── A7 ────────────────────────────────────────────────────────────────
    h2(doc, "A7. Le choix des socles SetFit")

    terme(doc, "Pourquoi pas de résultat mmBERT sur 20 Newsgroups",
          "Trois raisons, dans cet ordre. (1) L'atout de mmBERT est d'être multilingue ; sur un "
          "corpus intégralement anglophone, cet atout ne sert à rien. (2) L'angle « architecture "
          "moderne » y est déjà couvert par ModernBERT, qui est précisément la version "
          "anglophone de la même architecture — c'est la comparaison pertinente sur ce corpus. "
          "(3) mmBERT est environ cinq fois plus lent à entraîner : je l'ai réservé au cas "
          "d'usage métier, là où il apporte quelque chose.",
          "Formulation courte pour le jury : « Chaque corpus compare bien quatre modèles. La "
          "variante de socle est choisie selon la langue : ModernBERT sur l'anglais, mmBERT sur "
          "le français, mpnet servant de référence commune aux deux. »")


# ═══════════════════════════════════════════════════════════════════════════
SLIDES = [
    (1, "Couverture", "Se présenter en une phrase et annoncer le sujet.",
     ["« Jérémie Jambon, projet 9 : une preuve de concept sur la classification "
      "automatique des mails d'un club sportif. »",
      "Ne pas lire la diapositive. Enchaîner tout de suite."], []),

    (2, "Sommaire", "Annoncer les trois temps et leur durée : le jury sait où tu vas.",
     ["« Cinq minutes sur le plan prévisionnel, dix sur la démarche, cinq sur la "
      "démonstration du dashboard. »"], []),

    (3, "Introduction — le contexte", "Poser le problème métier AVANT toute technique.",
     ["Un logiciel pour clubs sportifs structurés ; sa brique centrale trie les mails entrants.",
      "Huit catégories métier.",
      "Insister sur LA difficulté : un club qui arrive n'a aucun mail déjà classé.",
      "C'est cette contrainte qui justifie tout le reste du projet."],
     [("Pourquoi huit catégories, et pas plus ou moins ?",
       "Elles correspondent aux flux réels d'un club structuré : inscriptions, sponsors, "
       "arbitrage, parents, fédération, logistique de match, indemnités, administratif. "
       "C'est le découpage qui permet de router un mail vers la bonne personne du bureau.")]),

    (4, "Les deux jeux de données", "Justifier d'avoir DEUX corpus, pas un.",
     ["Le corpus métier : 1 800 mails français, 225 par catégorie, parfaitement équilibré.",
      "Découpage 80/20 stratifié : 1 440 pour l'entraînement, 360 pour le test.",
      "Le corpus de validation : 20 Newsgroups, banc d'essai public reconnu.",
      "Le second sert à prouver que mes conclusions ne sont pas un artefact de mon corpus."],
     [("Pourquoi un corpus parfaitement équilibré, ce n'est pas réaliste ?",
       "C'est un choix méthodologique : l'équilibre isole l'effet du VOLUME de données, qui "
       "est ma variable d'étude. Si les classes étaient déséquilibrées, je ne saurais pas si "
       "un écart vient du peu d'exemples ou du déséquilibre. Le F1 macro et le class_weight "
       "équilibré préparent par ailleurs le cas déséquilibré réel."),
      ("Pourquoi 20 Newsgroups, un corpus des années 90 ?",
       "Justement parce qu'il est ancien et très utilisé : des centaines de publications y "
       "reportent leurs scores, je peux donc situer mes résultats. Et ses textes sont plus "
       "longs et plus bruités que des mails, ce qui teste la robustesse.")]),

    (5, "Construction du jeu de données", "Assumer le caractère synthétique et l'expliquer.",
     ["Aucun corpus public de mails de clubs sportifs n'existe : il a fallu le créer.",
      "Génération par grand modèle de langue, à partir de consignes structurées.",
      "Variations contrôlées : longueur, ton, profil d'expéditeur, dix sports différents.",
      "Argument RGPD : aucune donnée personnelle réelle n'est exposée."],
     [("Des données synthétiques, est-ce que ça ne fausse pas tout ?",
       "C'est ma première limite déclarée, et je l'assume. Le vocabulaire généré est "
       "probablement plus propre et plus séparable que des mails réels — c'est sans doute ce "
       "qui explique le 0,99 de la baseline en données complètes. Mais la QUESTION que je "
       "teste — quel modèle résiste au manque de données — reste valide, et je la vérifie "
       "sur un corpus réel, 20 Newsgroups, qui confirme la même hiérarchie."),
      ("Comment avez-vous garanti la qualité des étiquettes ?",
       "Chaque mail est généré pour une catégorie donnée, l'étiquette est donc native. J'ai "
       "ensuite vérifié la cohérence par l'analyse des mots les plus discriminants : ils "
       "correspondent au métier attendu pour chaque catégorie.")]),

    (6, "Les cinq stratégies", "Montrer que la comparaison est une logique produit, pas un catalogue.",
     ["Énumérer rapidement les cinq, sans les détailler — elles reviendront.",
      "Le message : chaque stratégie vise un profil de club différent.",
      "C'est ce qui mène à la matrice de décision en conclusion."],
     [("Pourquoi comparer autant de modèles pour un POC ?",
       "Parce que la vraie question n'est pas « quel est le meilleur modèle » mais « quel "
       "modèle pour quel client ». Un club sans historique et un club établi n'ont pas le "
       "même besoin. Les cinq stratégies couvrent quatre profils.")]),

    (7, "Pourquoi SetFit", "Justifier le choix par la littérature, pas par intuition.",
     ["Article fondateur : Tunstall et al. 2022, Hugging Face + Intel Labs + UKP Lab.",
      "Le chiffre choc : 8 exemples par classe rivalisent avec un fine-tuning sur 3 000.",
      "Sur RAFT : 71,3 % contre 62,7 % pour GPT-3, un modèle 30 fois plus gros.",
      "Entraînement en une trentaine de secondes, inférence locale gratuite."],
     [("Ces chiffres viennent de l'article — les avez-vous retrouvés ?",
       "Je n'ai pas reproduit RAFT, ce n'était pas l'objet. Mais j'ai retrouvé le PHÉNOMÈNE "
       "sur mes deux corpus : la supériorité en très faibles données. 0,95 contre 0,84 pour "
       "la baseline sur les mails, 0,79 contre 0,38 sur Newsgroups."),
      ("Pourquoi un algorithme de 2022 est-il « récent » ?",
       "Le critère est « moins de cinq ans, publié sur un support reconnu, disponible dans "
       "une librairie maintenue ». SetFit coche les trois, et j'ai poussé plus loin en "
       "testant des socles de 2024 et 2025.")]),

    (8, "Source 1 — Tunstall et al. (2022)",
     "L'article de recherche obligatoire. Montrer qu'on l'a vraiment lu.",
     ["Tunstall et al. 2022, publié conjointement par Hugging Face, Intel Labs et le UKP Lab "
      "de l'université technique de Darmstadt.",
      "Ce qu'il apporte : la méthode en deux étapes — comparaison de paires, puis "
      "classifieur léger.",
      "Le résultat de référence : 8 exemples par classe rivalisent avec un affinage sur "
      "3 000 exemples.",
      "Enchaîner vite : 20 à 25 secondes maximum."],
     [("Avez-vous lu l'article en entier ?",
       "Oui. Ce qui m'a le plus servi, c'est le mécanisme de génération de paires : c'est lui "
       "qui explique pourquoi SetFit résiste là où un affinage classique s'effondre. Je le "
       "réexplique en détail plus loin."),
      ("Que reproche l'article aux méthodes antérieures ?",
       "Leur dépendance à des consignes textuelles écrites à la main et à des modèles de "
       "plusieurs milliards de paramètres. PET, par exemple, transforme la classification en "
       "texte à trous : il faut rédiger la phrase-modèle et la table de correspondance "
       "mot-catégorie. SetFit supprime cette dépendance — c'est le sens de « prompt-free ».")]),

    (9, "Source 2 — Le blog Hugging Face (2022)",
     "La source du code réutilisé. C'est ta slide de transparence.",
     ["Rédigé par les auteurs de l'article eux-mêmes, avec du code exécutable.",
      "C'est de là que viennent la librairie setfit et la structure d'entraînement.",
      "Il détaille le banc d'essai RAFT : 71,3 % contre 62,7 % pour GPT-3."],
     [("Un article de blog, est-ce une source sérieuse ?",
       "Celui-ci l'est : il est écrit par les auteurs de l'article de recherche, publié sur le "
       "site officiel de Hugging Face, et il accompagne une librairie open source largement "
       "utilisée. Ce n'est pas un billet d'opinion, c'est la documentation de référence de la "
       "méthode."),
      ("Qu'avez-vous copié exactement ?",
       "La façon d'appeler la librairie : instancier le modèle, lancer l'entraînement. Le "
       "corpus, le protocole multi-régimes, les stratégies 4 et 5 et le tableau de bord sont "
       "de moi.")]),

    (10, "Source 3 — Ciancone et al. (2024)",
     "La source qui justifie un choix technique précis : le socle d'embeddings.",
     ["Premier banc d'essai massif de représentations de phrases en français.",
      "Une cinquantaine de modèles comparés, sur 8 familles de tâches.",
      "Sa conclusion : aucun modèle ne gagne partout, mais les bons multilingues rivalisent "
      "avec les modèles français spécialisés."],
     [("Pourquoi avoir besoin d'un banc d'essai pour choisir un socle ?",
       "Parce que SetFit ne part pas de zéro : il a besoin d'un modèle qui sait déjà mesurer "
       "la ressemblance entre phrases. Le choix de ce socle conditionne tout le reste, et "
       "MTEB-fr est le seul banc d'essai qui compare ces modèles sur du français."),
      ("Pourquoi ne pas avoir pris simplement le premier du classement ?",
       "Parce que le classement dépend de la tâche — c'est justement la conclusion de "
       "l'article : aucun modèle ne gagne partout. J'ai retenu un modèle multilingue bien "
       "classé et compatible avec mon environnement d'entraînement.")]),

    (11, "Les sources complémentaires",
     "Montrer que chaque brique technique a sa source. Slide à survoler si tu es en retard.",
     ["Reimers & Gurevych (2020) : la méthode de distillation qui a produit mon socle.",
      "Warner et al. (2024) : ModernBERT. Marone et al. (2025) : mmBERT.",
      "Martin et al. (2020) : CamemBERT, ma stratégie 2.",
      "Si le temps manque : « et quatre sources complémentaires pour chaque variante "
      "testée », puis enchaîner."],
     [("Comment définissez-vous un algorithme « récent » ?",
       "Trois critères : moins de cinq ans, publié sur un support reconnu, et disponible dans "
       "une librairie libre et maintenue. SetFit date de 2022, et j'ai poussé la comparaison "
       "jusqu'à des socles de 2024 et 2025.")]),

    (12, "SetFit : le principe", "Faire comprendre l'astuce contrastive à un non-spécialiste.",
     ["Étape 1 : on apprend au modèle à rapprocher les mails de même catégorie et à éloigner "
      "les autres. Étape 2 : une régression logistique légère sur ces représentations.",
      "L'astuce : 128 mails produisent plusieurs milliers de paires.",
      "La phrase à placer : « plutôt que d'apprendre — ce mail est une inscription — SetFit "
      "apprend d'abord — ces deux mails se ressemblent »."],
     [("Concrètement, combien de paires génère-t-on ?",
       "C'est quadratique : avec n exemples, on peut former de l'ordre de n² paires. Avec "
       "128 exemples, cela se compte en milliers. C'est justement pourquoi j'ai dû plafonner "
       "le nombre de pas d'entraînement sur les gros régimes."),
      ("Pourquoi une régression logistique en tête, et pas un réseau ?",
       "Parce qu'à ce stade le travail difficile est fait : les représentations séparent déjà "
       "bien les catégories. Une tête légère suffit, s'entraîne en quelques secondes, et "
       "reste interprétable.")]),

    (13, "Le système hybride", "Présenter l'architecture comme une décision d'ingénieur.",
     ["SetFit traite localement l'immense majorité des mails, gratuitement.",
      "Sous un seuil de confiance τ, le mail part vers le modèle génératif payant.",
      "τ est un curseur produit entre qualité, coût et latence.",
      "C'est ce qu'un éditeur logiciel déploierait vraiment."],
     [("Pourquoi vous fiez-vous à la confiance de SetFit et pas à celle du LLM ?",
       "Parce que la confiance de SetFit est une vraie probabilité, issue de sa régression "
       "logistique. Celle du LLM est auto-déclarée : il écrit un nombre dans son JSON, sans "
       "calcul derrière. Dans mes mesures elle valait très souvent 0,5, quelle que soit la "
       "situation."),
      ("Comment avez-vous choisi τ = 0,8 ?",
       "Par balayage de sept seuils, de 0,3 à 0,9. Et je le garde FIXE sur tous les régimes : "
       "choisir le meilleur seuil régime par régime reviendrait à l'ajuster sur le jeu de "
       "test, ce qui fausserait la comparaison.")]),

    (14, "Le protocole", "Prouver que la comparaison est méthodologiquement solide.",
     ["Cinq régimes : 8, 16, 32, 64 exemples par catégorie, puis données complètes.",
      "Cinq tirages par régime, pour mesurer la stabilité.",
      "Environ 380 expériences.",
      "Métrique : le F1 macro, pour qu'une petite catégorie pèse autant qu'une grande."],
     [("Pourquoi cinq seeds et pas une seule ?",
       "Parce qu'avec 8 exemples, le tirage compte énormément : tomber sur 8 mails très "
       "typiques ou 8 cas limites change tout. Cinq tirages donnent une moyenne et une "
       "dispersion — c'est la bande colorée autour de mes courbes."),
      ("Pourquoi le F1 macro plutôt que l'accuracy ?",
       "Exemple : 7 catégories à 100 mails et une à 10. Un modèle qui rate complètement la "
       "petite affiche 98,6 % d'accuracy, ce qui masque le problème. Son F1 macro tombe à "
       "0,875 et le révèle. En production, les catégories rares comme les indemnités sont "
       "justement celles qu'on ne peut pas se permettre d'ignorer.")]),

    (15, "Résultats sur les mails", "Le résultat central du projet.",
     ["SetFit atteint 0,95 dès 8 exemples par catégorie.",
      "CamemBERT s'effondre à 0,10.",
      "Les écarts se resserrent quand le volume augmente.",
      "Lecture de la courbe : plus une courbe est haute À GAUCHE, meilleur est le modèle "
      "quand les données manquent."],
     [("0,10 pour CamemBERT, ce n'est pas une erreur d'expérience ?",
       "Non, et j'ai vérifié. C'est une non-convergence : je garde 3 epochs sur tous les "
       "régimes pour que la courbe reste comparable. À 8 exemples par catégorie, cela ne fait "
       "que 12 pas d'optimisation — le modèle n'a pas eu le temps d'apprendre. Et 0,10 "
       "correspond au hasard sur 8 classes, ce qui confirme le diagnostic. Avec plus d'epochs "
       "à faible volume on remonterait ce point, mais la comparaison entre régimes deviendrait "
       "moins lisible."),
      ("Pourquoi TF-IDF fait-il déjà 0,84 à 8 exemples ?",
       "Parce que le vocabulaire métier est très discriminant : « vestiaires », « licence », "
       "« partenariat » suffisent presque à trancher. C'est aussi une conséquence du corpus "
       "synthétique, que j'assume comme limite.")]),

    (16, "Vérification sur 20 Newsgroups", "Montrer que le résultat n'est pas un artefact.",
     ["Même hiérarchie que sur les mails.",
      "Textes plus longs et plus bruités : c'est un test plus dur.",
      "ModernBERT n'apparaît que sur ce corpus, car il est anglophone."],
     [("Pourquoi mmBERT n'est-il pas sur ce graphique ?",
       "Parce que son atout est d'être multilingue, ce qui ne sert à rien sur un corpus "
       "anglophone. L'angle « architecture moderne » y est couvert par ModernBERT, qui en est "
       "la version anglaise. Et mmBERT est cinq fois plus lent : je l'ai réservé au cas "
       "métier. Chaque corpus compare bien quatre modèles."),
      ("Les scores sont plus bas ici, pourquoi ?",
       "Les textes sont longs, bruités, et les frontières entre catégories plus floues — "
       "distinguer talk.politics.guns de talk.religion.misc est intrinsèquement plus dur que "
       "distinguer une inscription d'une demande de sponsor.")]),

    (17, "Avec toutes les données", "Le retournement : la baseline suffit.",
     ["Sur les mails, les quatre approches se rejoignent autour de 0,99.",
      "Le vocabulaire métier est très discriminant.",
      "C'est ce qui justifie de déployer le modèle le plus léger."],
     [("Si tout se vaut à 0,99, à quoi sert SetFit ?",
       "À la situation qui pose problème : le démarrage. Un club établi n'a pas besoin de "
       "SetFit, un club qui arrive si. C'est tout l'objet de ma matrice de décision.")]),

    (18, "Le modèle génératif", "Montrer l'alternative sans entraînement.",
     ["Aucun entraînement : la tâche est décrite en français dans la consigne.",
      "0,79 sans aucun exemple, 0,90 avec 8 exemples dans la consigne.",
      "400 à 600 ms par mail, et un coût par appel.",
      "Intérêt : opérationnel immédiatement, sans une seule donnée étiquetée."],
     [("Pourquoi ne pas tout faire avec un LLM, alors ?",
       "Trois raisons : le coût est récurrent et proportionnel au volume, la latence est cent "
       "fois supérieure à un modèle local, et les données sortent de chez le client. SetFit "
       "tourne en local, gratuitement, une fois entraîné.")]),

    (19, "Le système hybride en chiffres", "Le meilleur résultat de l'étude.",
     ["À τ = 0,8 : F1 de 0,997, le meilleur score obtenu.",
      "Seuls 1,1 % des mails partent vers le modèle payant.",
      "Le gain est donc quasi gratuit.",
      "À placer si tu as le temps : à 8 exemples, le système escalade 13 % des mails contre "
      "1,1 % en données complètes — le coût s'auto-régule à mesure que le club mûrit."],
     [("0,997, n'est-ce pas du sur-apprentissage sur votre jeu de test ?",
       "C'est une question légitime. Le seuil τ est balayé sur le test, donc je le fixe à 0,8 "
       "pour tous les régimes plutôt que de prendre le meilleur seuil de chaque régime — ce "
       "qui serait de l'optimisation sur le test. En production, on calibrerait τ sur un jeu "
       "de validation séparé.")]),

    (20, "Le bug révélé par le LLM", "L'anecdote qui démontre ta rigueur. À raconter.",
     ["Premier constat : le modèle génératif plafonnait à 0,18, à peine mieux que le hasard.",
      "Vérification des réponses brutes : il classait en réalité correctement.",
      "La cause : scikit-learn trie les catégories par ordre alphabétique, indépendamment de "
      "l'ordre que j'avais demandé — mes étiquettes étaient décalées par rapport aux textes.",
      "Après correction : 0,82. Les modèles entraînés, eux, masquaient l'anomalie."],
     [("Pourquoi les modèles entraînés ne montraient-ils rien ?",
       "Parce qu'ils apprennent la correspondance qu'on leur donne, même si elle est fausse. "
       "Ils avaient appris « ce texte = étiquette 3 » de façon cohérente, donc leurs scores "
       "restaient bons. Seul un modèle NON entraîné juge sur le contenu réel — il sert de "
       "détecteur d'incohérence de la vérité-terrain."),
      ("Qu'est-ce que ça change à vos résultats ?",
       "Tous les résultats présentés sont post-correction. C'est aussi ce qui m'a fait "
       "revérifier l'ensemble de la chaîne d'étiquetage.")]),

    (21, "Interprétabilité globale", "Montrer que le modèle décide sur du sens métier.",
     ["La régression logistique est directement lisible : chaque mot porte un poids.",
      "Les mots décisifs correspondent au métier — c'est ce qui explique la robustesse.",
      "Citer deux exemples : « vestiaires » pour la logistique, « attestation » pour "
      "l'administratif."], []),

    (22, "Interprétabilité locale", "La convergence entre les deux familles de modèles.",
     ["Sur un même mail classé « logistique de match » : la régression logistique retient "
      "vestiaires, terrain, match, samedi.",
      "SetFit, expliqué par LIME, retient vestiaires, samedi, match, terrain.",
      "Deux familles très différentes s'appuient sur les mêmes indices."],
     [("Comment fonctionne LIME exactement ?",
       "On fabrique des centaines de variantes du mail en retirant des mots au hasard, on "
       "fait classer chacune par le modèle, et on regarde quels mots faisaient basculer la "
       "décision. On ajuste ensuite un petit modèle linéaire lisible autour de ce mail."),
      ("Cette explication est-elle fiable ?",
       "Partiellement, et je le déclare en limite : LIME est stochastique, le tirage des mots "
       "retirés étant aléatoire. Deux exécutions donnent des listes légèrement différentes. "
       "SHAP, fondé sur les valeurs de Shapley, offrirait des garanties théoriques — au prix "
       "d'un calcul bien plus lourd.")]),

    (23, "Synthèse des résultats", "Répondre frontalement à la question du projet.",
     ["SetFit est décisivement supérieur là où le problème existe : 0,95 contre 0,84.",
      "CamemBERT n'exprime son potentiel qu'avec un historique conséquent.",
      "La baseline reste imbattable en rapport performance / simplicité quand les données "
      "abondent.",
      "Le génératif permet le démarrage immédiat, l'hybride le meilleur compromis."], []),

    (24, "Le dashboard", "Passer en démonstration LIVE. Ne pas lire la diapositive.",
     ["Annoncer les trois volets, puis basculer sur le navigateur.",
      "La capture reste en secours si le réseau fait défaut.",
      "Justifier le modèle servi : TF-IDF, quelques mégaoctets, réponse immédiate, "
      "performance équivalente en données complètes.",
      "Montrer l'URL publique : c'est la preuve du déploiement cloud."],
     [("Pourquoi ne pas avoir déployé SetFit, votre meilleur modèle ?",
       "SetFit est meilleur en few-shot. Le dashboard démontre le scénario du club établi, et "
       "dans ce régime TF-IDF atteint 0,9917 contre 0,9916 : ils sont à égalité. Déployer un "
       "modèle 500 fois plus lourd pour un gain nul n'aurait pas de sens. La vraie limite, "
       "c'est que le dashboard ne montre qu'un profil de club sur les quatre.")]),

    (25, "Volets 2 et 3", "Démontrer, pas décrire. Ta capture d'écran est un atout : sers-t'en.",
     ["Tirer un mail au hasard, le classer, montrer la confiance et le bandeau vert de "
      "vérification.",
      "Puis basculer sur les courbes d'apprentissage.",
      "À exploiter absolument : sur ta capture, le modèle local annonce 80 % de confiance "
      "et voit juste, tandis que le modèle génératif trouve la même réponse mais annonce "
      "50 %. C'est la démonstration visuelle de ton argument sur la confiance non calibrée.",
      "Si le réseau lâche : la capture suffit, et les diapositives 15 à 19 contiennent les "
      "mêmes graphiques."],
     [("Cette confiance de 50 % du modèle génératif, qu'est-ce que ça veut dire ?",
       "Justement rien de fiable. Ce nombre n'est pas calculé : le modèle l'écrit dans sa "
       "réponse comme n'importe quel autre mot. Il annonçait très souvent 50 %, quelle que "
       "soit la difficulté du mail. La confiance du modèle local, elle, sort d'un vrai calcul "
       "de probabilité — c'est pourquoi mon système hybride se fonde sur elle.")]),

    (26, "Mise en production et accessibilité", "Montrer la maturité d'ingénierie.",
     ["Code versionné sur GitHub, application hébergée sur Streamlit Cloud.",
      "Dépendances allégées pour le déploiement.",
      "Clé d'API dans le mécanisme de secrets, jamais dans le code.",
      "Accessibilité : palette adaptée aux déficiences de vision des couleurs, aucune "
      "information portée par la couleur seule, navigation au clavier."],
     [("Qu'entendez-vous par « aucune information portée par la couleur seule » ?",
       "Sur mes graphiques, chaque courbe a une forme de marqueur distincte et une étiquette. "
       "Une personne daltonienne peut donc identifier chaque série sans distinguer les "
       "couleurs. C'est le critère 1.4.1 du référentiel WCAG.")]),

    (27, "Conclusion", "La matrice de décision : c'est ton livrable intellectuel.",
     ["Quatre profils de club, quatre stratégies.",
      "La phrase de fin : il n'existe pas de meilleur modèle universel, la bonne réponse "
      "dépend du volume de données et des contraintes de coût.",
      "Enchaîner sur les limites, sans les esquiver : corpus synthétique, dashboard qui ne "
      "démontre qu'un profil."],
     [("Si vous aviez trois mois de plus, que feriez-vous ?",
       "Trois choses, dans cet ordre. Valider sur des mails réels anonymisés, pour recalibrer "
       "les attentes. Servir SetFit depuis un service dédié, pour couvrir le scénario du club "
       "en installation. Et mesurer la calibration des scores de confiance, parce que tout le "
       "routage hybride en dépend."),
      ("Quelle est la principale faiblesse de votre travail ?",
       "Le corpus synthétique. Il est probablement plus propre et plus séparable que des mails "
       "réels, ce qui gonfle sans doute les scores absolus. Mais la comparaison ENTRE modèles "
       "reste valide, puisqu'ils affrontent tous exactement le même corpus, et la hiérarchie "
       "se confirme sur un corpus réel.")]),

    (28, "Merci", "Remercier, et inviter aux questions.", [], []),
]


def partie_b_slides(doc):
    doc.add_page_break()
    h1(doc, "PARTIE B — Slide par slide")
    para(doc, "Pour chaque diapositive : le message à faire passer, ce qu'il faut dire, "
              "et les questions que l'évaluateur peut poser à cet endroit précis.",
         italic=True, color=GRIS, size=10)

    for num, nom, message, dire, questions in SLIDES:
        h2(doc, f"Slide {num} — {nom}")
        champ(doc, "Message clé", message, couleur=BLEU)
        if dire:
            p = doc.add_paragraph()
            r = p.add_run("Ce qu'il faut dire :")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = VERT
            for d in dire:
                bullet(doc, d)
        for q, a in questions:
            qr(doc, q, a)


# ═══════════════════════════════════════════════════════════════════════════
def partie_c_transversal(doc):
    doc.add_page_break()
    h1(doc, "PARTIE C — Les questions transversales")
    para(doc, "Celles qui ne portent pas sur une diapositive mais sur l'ensemble du projet. "
              "Ce sont souvent les plus déstabilisantes.", italic=True, color=GRIS, size=10)

    transversales = [
        ("En une phrase, votre algorithme récent est-il plus performant que la baseline ?",
         "Oui, là où le problème se pose : 0,95 contre 0,84 à 8 exemples par catégorie sur les "
         "mails, 0,79 contre 0,38 sur 20 Newsgroups. L'écart disparaît en données complètes, "
         "non parce que SetFit régresse, mais parce que la tâche cesse d'être difficile."),
        ("Qu'est-ce qui est vraiment de vous dans ce projet ?",
         "La librairie SetFit et la structure de pipeline viennent du tutoriel officiel, je le "
         "déclare. Tout le reste est original : la création du corpus métier, le protocole "
         "multi-régimes à cinq tirages, les stratégies 4 et 5 — le LLM comme classifieur et le "
         "système hybride —, l'analyse d'interprétabilité et le dashboard."),
        ("Votre POC est-il industrialisable en l'état ?",
         "Non, et ce n'est pas son objet. C'est une preuve de concept : elle démontre qu'une "
         "stratégie existe pour chaque profil de client. Pour industrialiser, il faudrait "
         "valider sur des mails réels, servir SetFit depuis un service dédié, calibrer les "
         "scores de confiance et mettre en place un suivi de dérive."),
        ("Combien ça coûte de faire tourner ça pour un club ?",
         "Le modèle déployé est gratuit à l'inférence : c'est du calcul local. Si on active "
         "l'hybride, on paie environ 0,000022 $ par mail escaladé — soit à peine plus de deux "
         "centimes pour mille mails escaladés. Et l'escalade diminue à mesure que le club "
         "accumule des données."),
        ("Comment savez-vous que vos résultats sont fiables ?",
         "Trois garde-fous. Cinq tirages par régime, qui donnent la dispersion visible sur mes "
         "courbes. Un second corpus, public, qui confirme la même hiérarchie. Et un "
         "échantillonnage stratifié reproductible, donc n'importe qui peut relancer mon code "
         "et retrouver exactement les mêmes chiffres."),
        ("Et la conformité RGPD ?",
         "Trois points. Mon corpus est synthétique, aucune donnée personnelle réelle n'est "
         "exposée. Le modèle déployé tourne en local, les mails ne sortent pas. Et l'API "
         "utilisée pour l'hybride est celle d'un fournisseur européen, hébergée en Europe — "
         "ce qui compte si des mails réels devaient y transiter."),
        ("Pourquoi avoir abandonné le modèle français initialement prévu ?",
         "Le modèle français spécialisé envisagé au départ présentait une incompatibilité avec "
         "mon environnement d'entraînement. Je l'ai remplacé par un socle multilingue de "
         "référence du banc d'essai MTEB-fr, ce qui est cohérent avec la conclusion de ce même "
         "banc d'essai : les bons modèles multilingues rivalisent avec les modèles spécialisés."),
        ("Qu'avez-vous appris que vous ne saviez pas au départ ?",
         "Deux choses. D'abord qu'un modèle plus gros et plus récent ne gagne pas "
         "automatiquement : mmBERT, le plus moderne, n'a pas dépassé un socle de 2020 pour "
         "cinq fois plus de calcul. Ensuite qu'un modèle non entraîné peut servir d'outil de "
         "contrôle qualité — c'est le LLM qui a révélé un bug dans ma vérité-terrain."),
    ]
    for q, a in transversales:
        qr(doc, q, a)

    # ── Antisèche ────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "L'antisèche — les chiffres à connaître par cœur")

    h2(doc, "Mes résultats (mails de clubs sportifs)")
    for t in [
        "8 exemples/catégorie : SetFit 0,95 | baseline 0,84 | CamemBERT 0,10",
        "Données complètes : baseline 0,9917 | SetFit 0,9916 — égalité stricte",
        "Ministral : 0,79 sans exemple, 0,90 avec 8 exemples dans la consigne",
        "Hybride à τ=0,8 : 0,95 en 8-shot, 0,997 en données complètes",
        "Escalade hybride : 13 % à 8 exemples, 1,1 % en données complètes",
    ]:
        bullet(doc, t)

    h2(doc, "Mes résultats (20 Newsgroups)")
    for t in [
        "8 exemples : SetFit 0,79 | baseline 0,38 | CamemBERT 0,11",
        "Données complètes : SetFit 0,84 | baseline 0,80 | ModernBERT 0,80 | CamemBERT 0,79",
        "Ministral : 0,82 — après correction du bug d'étiquetage (0,18 avant)",
    ]:
        bullet(doc, t)

    h2(doc, "La littérature")
    for t in [
        "SetFit : 8 exemples/classe rivalisent avec un fine-tuning sur 3 000 exemples",
        "RAFT : SetFit 355 M = 71,3 % contre GPT-3 175 B = 62,7 %",
        "Entraînement SetFit : ~30 s pour 0,025 $",
    ]:
        bullet(doc, t)

    h2(doc, "Le dispositif")
    for t in [
        "1 800 mails, 225 par catégorie, 8 catégories, découpage 80/20 stratifié",
        "5 régimes × 5 tirages ≈ 380 expériences",
        "Modèle déployé : 2,3 Mo, environ 500 fois plus léger que SetFit",
    ]:
        bullet(doc, t)


SCRIPTS = [
    ("generate_dataset.py", "Crée le corpus métier",
     "Génère les 1 800 mails via un grand modèle de langue, à partir de consignes "
     "structurées. Reprend automatiquement là où il s'est arrêté en cas d'interruption : "
     "les catégories déjà produites sont détectées et sautées."),
    ("run_baseline.py", "Stratégie 1 — TF-IDF + régression logistique",
     "Lance 25 expériences par corpus (5 régimes × 5 tirages) et enregistre chaque "
     "résultat dans un fichier séparé. C'est la référence à battre."),
    ("run_camembert.py", "Stratégie 2 — affinage de CamemBERT",
     "Mêmes 25 expériences par corpus. Chaque run réentraîne CamemBERT depuis le modèle "
     "pré-entraîné sur le sous-ensemble correspondant. C'est le script le plus long : "
     "quelques minutes par run sur données complètes."),
    ("run_setfit.py", "Stratégie 3 — SetFit, l'algorithme récent étudié",
     "Lance les runs pour les trois socles, sur les corpus compatibles avec chacun : "
     "ModernBERT étant anglophone et mmBERT multilingue, tous ne tournent pas sur les "
     "deux corpus."),
    ("run_ministral.py", "Stratégie 4 — le modèle génératif comme classifieur",
     "Aucun entraînement : évalue directement le modèle via l'API, en deux modes — sans "
     "aucun exemple, puis avec 8 exemples glissés dans la consigne."),
    ("run_hybrid.py", "Stratégie 5 — le routage par confiance",
     "Pour chaque configuration, entraîne SetFit puis balaie sept seuils de confiance τ, "
     "en escaladant vers le modèle génératif les mails sous le seuil. Produit le couple "
     "qualité / taux d'escalade pour chaque τ."),
    ("run_all.py", "L'orchestrateur",
     "Enchaîne les six étapes dans l'ordre logique, chaque script tournant dans son propre "
     "processus — ce qui isole la mémoire et rend les erreurs traçables. C'est le point "
     "d'entrée pour tout reproduire."),
    ("dump_predictions.py", "Prépare l'analyse qualitative des erreurs",
     "Entraîne chaque modèle sur une configuration de référence, puis enregistre pour "
     "chaque mail de test la vraie catégorie et la prédiction de chaque modèle. Alimente "
     "le notebook d'analyse des erreurs."),
    ("feature_importance.py", "Produit l'analyse d'interprétabilité",
     "Calcule l'importance globale (les mots de plus fort poids par catégorie, lisibles "
     "directement dans la régression logistique) et l'importance locale via LIME pour "
     "SetFit, qui est une boîte noire."),
    ("save_deploy_model.py", "Fabrique le modèle mis en ligne",
     "Entraîne et enregistre le modèle TF-IDF léger servi par le tableau de bord. Les "
     "modèles lourds pèsent plusieurs gigaoctets et réclament une carte graphique : "
     "impossibles à héberger sur l'offre gratuite."),
    ("make_slide_figures.py", "Génère les graphiques du support",
     "Produit les figures à partir des résultats réels : courbes d'apprentissage, "
     "comparaison sur données complètes, arbitrage du système hybride. Palette adaptée "
     "aux déficiences de vision des couleurs, doublée d'un encodage par forme."),
]


DASHBOARD = [
    ("dashboard/app.py", "L'application elle-même (environ 520 lignes)",
     "Le point d'entrée du tableau de bord : il construit les trois volets, charge le "
     "modèle léger et les résultats d'expériences, et gère l'appel optionnel au modèle "
     "génératif. C'est ce fichier que Streamlit Cloud exécute."),
    ("dashboard/text_analysis.py", "Les analyses de l'onglet exploration",
     "Regroupe les calculs sur le texte : longueur des messages, fréquence des mots, "
     "nuage de mots. Séparé de l'application pour que la logique d'analyse reste testable "
     "indépendamment de l'affichage."),
    ("dashboard/theme.py", "La charte visuelle et l'accessibilité",
     "Centralise la palette et les repères de mise en forme. C'est ici qu'est appliquée "
     "la conformité au référentiel WCAG : palette distinguable en cas de déficience de "
     "vision des couleurs, et information jamais portée par la couleur seule."),
]


def partie_d_scripts(doc):
    doc.add_page_break()
    h1(doc, "PARTIE D — À quoi sert chaque fichier")
    para(doc, "De quoi répondre si le jury demande comment le projet est organisé, ou "
              "comment reproduire les résultats.", italic=True, color=GRIS, size=10)

    h2(doc, "L'application déployée")
    for nom, role, detail in DASHBOARD:
        p = doc.add_paragraph()
        r = p.add_run(nom)
        r.bold = True
        r.font.color.rgb = BLEU
        r.font.size = Pt(11)
        r2 = p.add_run(f"  —  {role}")
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = VERT
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p2 = doc.add_paragraph()
        r3 = p2.add_run(detail)
        r3.font.size = Pt(10.5)
        p2.paragraph_format.space_after = Pt(4)

    para(doc, "Pour la lancer en local : streamlit run dashboard/app.py",
         italic=True, color=GRIS, size=10)

    h2(doc, "Les scripts d'expérimentation")
    for nom, role, detail in SCRIPTS:
        p = doc.add_paragraph()
        r = p.add_run(nom)
        r.bold = True
        r.font.color.rgb = BLEU
        r.font.size = Pt(11)
        r2 = p.add_run(f"  —  {role}")
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = VERT
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)

        p2 = doc.add_paragraph()
        r3 = p2.add_run(detail)
        r3.font.size = Pt(10.5)
        p2.paragraph_format.space_after = Pt(4)

    para(doc, "Pour tout reproduire : une seule commande, python scripts/run_all.py — "
              "elle enchaîne la génération du corpus puis les cinq stratégies.",
         italic=True, color=GRIS, size=10)

    h2(doc, "Comment le code est rangé")
    for t, d in [
        ("src/data/", "chargement des deux corpus et tirage stratifié reproductible"),
        ("src/models/", "une implémentation par stratégie (baseline, CamemBERT, SetFit, "
                        "Ministral, hybride)"),
        ("src/evaluation/", "calcul des métriques et agrégation des résultats"),
        ("scripts/", "les points d'entrée qui lancent les expériences"),
        ("dashboard/", "l'application, qui importe src/ pour charger données et modèles"),
        ("notebooks/", "les quatre analyses : exploration, résultats, erreurs, décision"),
        ("results/", "les 383 résultats d'expériences et leur agrégat"),
    ]:
        bullet(doc, d, f"{t} — ")

    para(doc, "Le principe : src/ contient la logique réutilisable, scripts/ et dashboard/ "
              "ne font que l'appeler. C'est ce qui permet au tableau de bord et aux "
              "notebooks de partager exactement le même code de chargement des données.",
         italic=True, color=GRIS, size=10)


def build():
    doc = Document()
    titre(doc, "Note de révision — Soutenance")
    para(doc, "POC Classification de mails de clubs sportifs — Projet 9",
         italic=True, color=GRIS, size=12)
    para(doc, "Quatre parties : le glossaire de tous les termes techniques, le déroulé "
              "diapositive par diapositive avec les questions probables, les questions "
              "transversales suivies de l'antisèche chiffrée, et le rôle de chaque script "
              "du projet.", color=GRIS, size=10)

    p = doc.add_paragraph()
    r = p.add_run("28 diapositives — 5 min (plan prévisionnel) · 10 min (démarche) · "
                  "5 min (dashboard)")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = VERT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "Rythme à tenir : diapositives 4 à 11 en 5 minutes, soit environ 35 secondes "
              "chacune — les quatre diapositives de sources doivent passer en 20 à 25 secondes. "
              "Diapositives 12 à 23 en 10 minutes, soit 50 secondes. Diapositives 24 à 26 en "
              "démonstration live.", italic=True, color=GRIS, size=10)

    doc.add_page_break()
    partie_a_glossaire(doc)
    partie_b_slides(doc)
    partie_c_transversal(doc)
    partie_d_scripts(doc)

    doc.save(OUT)
    print(f"Document généré : {OUT}")


if __name__ == "__main__":
    build()
