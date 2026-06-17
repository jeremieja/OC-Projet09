"""
Génère un document Word de préparation à la soutenance : synthèse des références
bibliographiques du POC, orientée "ce que je dois retenir et savoir défendre".

Sortie : Desktop/Projet 9/Preparation_soutenance_biblio.docx

Usage :
    python scripts/generate_biblio_doc.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parents[1] / "Preparation_soutenance_biblio.docx"

# Couleurs
BLEU = RGBColor(0x0F, 0x4C, 0x81)
VERT = RGBColor(0x1B, 0x7F, 0x4B)
GRIS = RGBColor(0x55, 0x55, 0x55)


def h_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = BLEU
    return p


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = BLEU
    r.font.size = Pt(16)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    r.font.size = Pt(13)
    return p


def para(doc, text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def field(doc, label, text, label_color=VERT):
    """Ligne 'Label : contenu' avec label en couleur/gras."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label} : ")
    r.bold = True
    r.font.color.rgb = label_color
    p.add_run(text)
    return p


def build():
    doc = Document()

    # ── Page de garde ────────────────────────────────────────────────
    h_title(doc, "Préparation soutenance — Synthèse bibliographique")
    para(doc,
         "POC Classification de mails de clubs sportifs — DataSpace, Projet 9",
         italic=True, color=GRIS, size=12)
    para(doc,
         "Document de révision : pour chaque référence, l'essentiel à retenir, "
         "les chiffres à citer, son rôle dans le POC, et les questions probables du jury.",
         color=GRIS, size=10)
    doc.add_paragraph()

    # ── Antisèche express ────────────────────────────────────────────
    h1(doc, "Antisèche express (chiffres à connaître par cœur)")
    bullet(doc, "SetFit = 8 exemples/classe ≈ RoBERTa Large entraîné sur 3 000 exemples.",
           "Few-shot : ")
    bullet(doc, "SetFit 355M params bat GPT-3 175B sur RAFT (71,3 % vs 62,7 %).",
           "Efficacité : ")
    bullet(doc, "SetFit s'entraîne en 30 s / 0,025 $ (vs 11 min / 0,70 $ pour T-Few 3B).",
           "Coût : ")
    bullet(doc, "CamemBERT s'effondre en few-shot (F1 0,10 à 8-shot) puis atteint 0,99 en full data.",
           "Mes résultats : ")
    bullet(doc, "SetFit champion few-shot (0,95 sur emails à 8-shot) ; hybride 0,997 en escaladant ~1 % des mails.",
           "Mes résultats : ")
    bullet(doc, "Ministral zero-shot : 0,79 (emails, labels parlants) mais 0,18 puis 0,82 sur Newsgroups APRÈS correction d'un bug de labels.",
           "Mes résultats : ")

    doc.add_page_break()

    # ── Les références ───────────────────────────────────────────────
    h1(doc, "Les références, une par une")

    refs = [
        {
            "titre": "1. Tunstall et al. (2022) — SetFit : Efficient Few-Shot Learning Without Prompts",
            "ref": "arXiv:2209.11055 — article de recherche fondateur (référence obligatoire).",
            "role": "Cœur de ma stratégie 3. C'est LA méthode few-shot du POC, celle qui "
                    "résout le problème du « nouveau club sans données étiquetées ».",
            "retenir": [
                "Problème résolu : les méthodes few-shot d'avant (PET, PEFT) dépendent de "
                "prompts manuels fragiles et de modèles à plusieurs milliards de paramètres.",
                "Méthode en 2 étapes : (1) fine-tuning CONTRASTIF d'un Sentence Transformer "
                "sur des paires (même classe = positif / classes ≠ = négatif) ; (2) une tête "
                "de classification légère (régression logistique) sur les embeddings produits.",
                "Astuce clé : la génération de paires démultiplie le signal — 16 exemples × 8 "
                "classes = 128 exemples → plusieurs milliers de paires d'entraînement.",
                "Sans prompt ni verbalizer, et 'plusieurs ordres de grandeur' moins de paramètres.",
            ],
            "chiffres": [
                "8 exemples/classe sur Customer Reviews ≈ RoBERTa Large fine-tuné sur 3 000 exemples.",
                "RAFT : SetFit (355M) = 71,3 % > PET (69,6 %) > GPT-3 175B (62,7 %) ; T-Few 11B = 75,8 %.",
                "Bat la performance humaine sur 7 des 11 tâches RAFT.",
                "Entraînement : 30 s sur V100 pour 0,025 $ (vs T-Few 3B : 11 min / 0,70 $) → 28× plus rapide.",
            ],
            "jury": "« Pourquoi SetFit et pas un simple fine-tuning ? » → Parce qu'avec 8-30 "
                    "exemples par classe (cas du nouveau club), le fine-tuning classique "
                    "(CamemBERT) s'effondre (mon F1=0,10), alors que l'entraînement contrastif "
                    "de SetFit exploite des milliers de paires et atteint 0,95.",
        },
        {
            "titre": "2. Hugging Face Blog (2022) — SetFit (article de vulgarisation officiel)",
            "ref": "huggingface.co/blog/setfit — illustre le papier avec du code de référence.",
            "role": "Source du code de pipeline que j'ai réutilisé (déclaration de transparence) "
                    "et des chiffres de comparaison que je cite.",
            "retenir": [
                "Présente l'API setfit (librairie Apache 2.0) et la structure d'entraînement.",
                "Donne le tableau RAFT détaillé et les coûts d'entraînement comparés.",
                "Mentionne la distillation : jusqu'à 123× d'accélération à l'inférence.",
            ],
            "chiffres": [
                "Tableau RAFT complet (cf. référence 1).",
                "Distillation → inférence 123× plus rapide.",
            ],
            "jury": "« Qu'avez-vous réutilisé ? » → La librairie setfit et la structure de "
                    "pipeline du tutoriel ; tout le reste (datasets, protocole multi-régimes, "
                    "niveaux 4-5, dashboard) est original.",
        },
        {
            "titre": "3. Martin et al. (2020) — CamemBERT",
            "ref": "arXiv:1911.03894 — ACL 2020. Modèle de référence pour le français.",
            "role": "Ma stratégie 2 (fine-tuning classique). Représente l'approche "
                    "« data-driven » : excellent SI le club a beaucoup de données.",
            "retenir": [
                "Premier grand modèle BERT monolingue français, entraîné sur des données web "
                "(OSCAR) plutôt que Wikipédia.",
                "Découverte marquante : 4 Go de web bien choisi = aussi bon que 130+ Go.",
                "Atteint ou dépasse l'état de l'art sur 4 tâches FR : POS, parsing, NER, NLI.",
            ],
            "chiffres": [
                "Corpus de seulement 4 Go suffisant (vs 130+ Go).",
                "État de l'art sur les 4 tâches FR évaluées.",
            ],
            "jury": "« CamemBERT est censé être le meilleur en français, pourquoi est-il "
                    "dernier en few-shot ? » → Parce qu'il est 'data-hungry' : un gros réseau "
                    "(110M params) ne peut pas apprendre avec 8 exemples. Mon expérience le "
                    "montre : F1=0,10 à 8-shot, mais 0,99 en full data.",
        },
        {
            "titre": "4. Ciancone et al. (2024) — MTEB-fr",
            "ref": "arXiv:2405.20468 — extension du Massive Text Embedding Benchmark au français.",
            "role": "Justifie scientifiquement mon choix de backbone pour SetFit et de "
                    "CamemBERT. Benchmark de référence pour choisir un modèle d'embedding FR.",
            "retenir": [
                "Premier benchmark massif d'embeddings de phrases en français.",
                "15 datasets existants + 3 nouveaux, 8 catégories de tâches, 51 modèles comparés.",
                "Conclusion clé : aucun modèle ne gagne partout, MAIS les gros modèles "
                "MULTILINGUES pré-entraînés sur la similarité de phrases sont excellents.",
            ],
            "chiffres": [
                "51 modèles évalués, 8 catégories de tâches, 18 datasets.",
            ],
            "jury": "« Pourquoi un backbone multilingue (mpnet) et pas un modèle français "
                    "spécialisé ? » → MTEB-fr montre que les bons modèles multilingues sur "
                    "similarité de phrases rivalisent voire dépassent les spécialisés FR ; "
                    "et techniquement, sentence-camembert-large crashait mon GPU Blackwell.",
        },
        {
            "titre": "5. Reimers & Gurevych (2020) — Sentence Embeddings multilingues par distillation",
            "ref": "arXiv:2004.09813 — EMNLP 2020.",
            "role": "Fondement technique de mon backbone SetFit retenu "
                    "(paraphrase-multilingual-mpnet-base-v2).",
            "retenir": [
                "Méthode pour rendre multilingue un modèle d'embedding anglais par "
                "distillation de connaissances (un modèle 'élève' multilingue imite un "
                "'professeur' anglais).",
                "Permet d'aligner les embeddings de plusieurs langues dans le même espace.",
            ],
            "chiffres": [
                "Backbone retenu : paraphrase-multilingual-mpnet-base-v2 (issu de ces travaux).",
            ],
            "jury": "« D'où vient votre backbone ? » → C'est un mpnet rendu multilingue par "
                    "distillation (cette méthode), parfaitement adapté à des mails français.",
        },
        {
            "titre": "6. Warner et al. (2024) — ModernBERT",
            "ref": "arXiv:2412.13663 — Answer.AI & LightOn. Encodeur de référence 2024.",
            "role": "Variante de backbone SetFit que je teste pour la question centrale : "
                    "modernité architecturale vs modèle multilingue établi.",
            "retenir": [
                "Première vraie modernisation de l'architecture BERT depuis l'original.",
                "Contexte natif de 8 192 tokens (vs 512 pour BERT/CamemBERT).",
                "Pensé pour l'inférence rapide sur GPU courants : meilleure efficacité mémoire.",
            ],
            "chiffres": [
                "Contexte 8 192 tokens (16× plus que BERT).",
                "Entraîné sur 2 000 milliards de tokens.",
                "État de l'art sur classification + retrieval, encodeur le plus rapide/efficace.",
            ],
            "jury": "« Pourquoi tester ModernBERT ? » → Pour voir si une archi récente (2024) "
                    "compense l'absence de spécialisation FR. Sur Newsgroups (EN) il est bon "
                    "(0,80 full) mais reste derrière mpnet multilingue (0,84).",
        },
        {
            "titre": "7. Wasserblat (2025) — SetFit + ModernBERT en few-shot",
            "ref": "Medium — évaluation récente de l'association SetFit + ModernBERT.",
            "role": "Justifie l'exploration de backbones ModernBERT/mmBERT dans mon SetFit.",
            "retenir": [
                "Montre que ModernBERT comme backbone SetFit excelle en few-shot sur des "
                "textes longs (IMDB, ArXiv).",
                "Confirme que l'astuce SetFit (contrastif) reste valable avec des encodeurs récents.",
            ],
            "chiffres": [
                "IMDB : 92,7 % avec 8 exemples/classe (vs 62,5 % pour ModernBERT seul) — +50 %.",
                "ArXiv-new : 90,3 % avec 64 exemples (vs 76,5 % pour ModernBERT seul).",
            ],
            "jury": "« Votre idée de tester mmBERT vient d'où ? » → De ce travail qui valide "
                    "SetFit+ModernBERT ; je l'ai prolongé avec mmBERT (ModernBERT multilingue, 2025).",
        },
        {
            "titre": "8. Marone et al. (2025) — mmBERT",
            "ref": "huggingface.co/blog/mmbert — JHU-CLSP. Le backbone le plus récent testé.",
            "role": "Variante de backbone SetFit la plus récente : le seul à réunir modernité "
                    "architecturale ET couverture multilingue native. Tranche ma question centrale.",
            "retenir": [
                "ModernBERT rendu multilingue : 1 800+ langues, 3 000 milliards de tokens.",
                "Entraînement en 3 phases avec masquage décroissant (30 % → 15 % → 5 %) et ajout "
                "progressif des langues (60 → 110 → 1 833).",
                "Égale les modèles 100 % anglais sur MTEB-EN avec <25 % de données anglaises.",
            ],
            "chiffres": [
                "1 800+ langues, 3T+ tokens.",
                "mmBERT-base : 307M params (110M hors embeddings) ; small : 140M.",
                "2 à 4× plus rapide que les multilingues précédents ; contexte 8 192 tokens.",
            ],
            "jury": "« mmBERT a-t-il gagné ? » → Sur mes emails il est très bon (~0,99 full) "
                    "mais à égalité avec mpnet, pour un coût de calcul ~5× supérieur. "
                    "Conclusion : la modernité archi ne suffit pas à justifier le surcoût ici.",
        },
        {
            "titre": "9. Mistral AI (2025) — Ministral 8B (documentation La Plateforme)",
            "ref": "docs.mistral.ai — modèles edge Apache 2.0, hébergement européen (RGPD).",
            "role": "Ma stratégie 4 (LLM-as-classifier) et brique du système hybride (stratégie 5).",
            "retenir": [
                "LLM utilisé sans entraînement, via API : on lui DEMANDE de classer en langage naturel.",
                "Hébergement européen, conformité RGPD native — argument produit important.",
                "Deux modes : zero-shot (instruction seule) et few-shot (8 exemples in-context).",
            ],
            "chiffres": [
                "Tarif : 0,10 $ / million de tokens (entrée et sortie).",
                "Latence ~400-600 ms par mail dans mes mesures.",
                "Mes F1 : emails zero-shot 0,79, few-shot 0,90 ; Newsgroups 0,82 (après fix labels).",
            ],
            "jury": "« Pourquoi Ministral était nul (0,18) sur Newsgroups au début ? » → "
                    "C'était un BUG de ma vérité-terrain (labels décalés vs textes). Le LLM "
                    "classait juste sur le vrai contenu ; c'est lui qui a RÉVÉLÉ le bug. "
                    "Après correction : 0,82. Belle leçon de rigueur.",
        },
    ]

    for r in refs:
        h2(doc, r["titre"])
        para(doc, r["ref"], italic=True, color=GRIS, size=10)
        field(doc, "Rôle dans mon POC", r["role"])

        p = doc.add_paragraph()
        run = p.add_run("À retenir :")
        run.bold = True
        run.font.color.rgb = BLEU
        for item in r["retenir"]:
            bullet(doc, item)

        p = doc.add_paragraph()
        run = p.add_run("Chiffres à citer :")
        run.bold = True
        run.font.color.rgb = VERT
        for item in r["chiffres"]:
            bullet(doc, item)

        field(doc, "Question probable du jury", r["jury"], label_color=RGBColor(0xB8, 0x5C, 0x00))
        doc.add_paragraph()

    # ── Fil narratif final ───────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Le fil narratif à dérouler en soutenance")
    steps = [
        ("Le problème métier", "Un SaaS pour clubs sportifs doit trier les mails entrants. "
         "Difficulté : un nouveau club arrive avec ZÉRO mail étiqueté."),
        ("La démarche", "Comparer 5 stratégies sur une courbe d'apprentissage (8 → full data), "
         "sur 2 datasets (emails FR + 20 Newsgroups pour la validation scientifique)."),
        ("Le résultat clé n°1", "SetFit est le champion du few-shot : 0,95 dès 8 exemples/classe. "
         "C'est la réponse au nouveau club."),
        ("Le résultat clé n°2", "CamemBERT est data-hungry : nul en few-shot (0,10), excellent en "
         "full data (0,99). À réserver aux clubs établis."),
        ("Le résultat clé n°3", "Le système hybride atteint 0,997 en n'escaladant que ~1 % des mails "
         "vers le LLM payant : le seuil τ devient un levier produit qualité/coût."),
        ("L'anecdote qui marque", "Ministral semblait nul sur Newsgroups (0,18). En creusant, j'ai "
         "découvert un bug de labels dans ma vérité-terrain. Le LLM avait raison ! Après "
         "correction : 0,82. La rigueur paie."),
        ("La livraison", "Matrice de décision produit + dashboard Streamlit déployé sur le cloud "
         "(modèle TF-IDF léger, égal à SetFit en full data mais 500× plus petit)."),
    ]
    for i, (t, txt) in enumerate(steps, 1):
        bullet(doc, txt, bold_prefix=f"{t} — ")

    doc.save(OUT)
    print(f"Document généré : {OUT}")


if __name__ == "__main__":
    build()
