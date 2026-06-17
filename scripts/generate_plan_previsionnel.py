"""
Génère le Plan Prévisionnel condensé (.docx), conforme au modèle imposé (1 page) :
  - Dataset retenu (quelques lignes)
  - Modèle envisagé (arguments justifiant le choix + objectif/contexte)
  - Références bibliographiques (2-3, dont 1 article de recherche obligatoire)
  - Démarche de test (baseline + méthode récente)

Sortie : Desktop/Projet 9/Plan_previsionnel.docx
Usage : python scripts/generate_plan_previsionnel.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parents[1] / "Plan_previsionnel.docx"
BLEU = RGBColor(0x0F, 0x4C, 0x81)
GRIS = RGBColor(0x55, 0x55, 0x55)


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = BLEU


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLEU


def body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(size)
    return p


def bullet(doc, text, bold_prefix=None, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(size)
    r = p.add_run(text); r.font.size = Pt(size)
    return p


def build():
    doc = Document()
    # Marges serrées pour tenir en 1 page
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    title(doc, "Plan prévisionnel — POC Classification de mails de clubs sportifs")

    # ── Dataset ───────────────────────────────────────────────────────
    h1(doc, "Dataset retenu")
    body(doc,
         "Deux jeux de données complémentaires. (1) Dataset métier : ~1 800 mails de clubs "
         "sportifs en français, générés via LLM et étiquetés sur 8 catégories (inscription, "
         "sponsor, arbitrage-officiels, parent, fédération, logistique-matchday, "
         "indemnités-coachs, divers-administratif). (2) Dataset de validation : sous-ensemble "
         "équilibré de 8 classes de 20 Newsgroups, benchmark public, pour la reproductibilité "
         "scientifique. Enjeu central : un nouveau club arrive avec zéro mail étiqueté.")

    # ── Modèle envisagé ───────────────────────────────────────────────
    h1(doc, "Modèle envisagé")
    body(doc,
         "Algorithme récent étudié : SetFit (Tunstall et al., 2022), framework de "
         "fine-tuning few-shot et prompt-free. Objectif : classer des mails avec très peu "
         "d'exemples étiquetés (~10-30 par catégorie), sans GPU obligatoire ni dépendance "
         "externe — exactement la contrainte d'un club sans historique.")
    body(doc, "Arguments justifiant le choix (résultats issus des articles de référence) :")
    bullet(doc, "Avec 8 exemples/classe, SetFit est compétitif avec un fine-tuning de "
                "RoBERTa Large sur 3 000 exemples.")
    bullet(doc, "Sur le benchmark RAFT, SetFit (355M paramètres) atteint 71,3 %, surpassant "
                "GPT-3 175B (62,7 %) — un modèle 30× plus petit.")
    bullet(doc, "Entraînement en ~30 s pour 0,025 $, déployable localement et gratuit à "
                "l'inférence.")
    body(doc,
         "Backbone retenu : paraphrase-multilingual-mpnet-base-v2 (référence MTEB-fr). "
         "Variantes testées pour évaluer la modernité architecturale : ModernBERT (2024) et "
         "mmBERT (2025, ModernBERT multilingue).")

    # ── Références ────────────────────────────────────────────────────
    h1(doc, "Références bibliographiques")
    bullet(doc, "Tunstall et al. (2022), Efficient Few-Shot Learning Without Prompts, "
                "arXiv:2209.11055 — article de recherche fondateur de SetFit.",
                bold_prefix="[Recherche] ")
    bullet(doc, "Hugging Face Blog (2022), SetFit, huggingface.co/blog/setfit — illustration "
                "appliquée du papier, avec code de référence.",
                bold_prefix="[Blog] ")
    bullet(doc, "Ciancone et al. (2024), Extending MTEB to French, arXiv:2405.20468 — "
                "benchmark justifiant le choix du backbone d'embedding.",
                bold_prefix="[Recherche] ")

    # ── Démarche POC ──────────────────────────────────────────────────
    h1(doc, "Démarche de test (preuve de concept)")
    body(doc,
         "Comparaison multi-régimes (8, 16, 32, 64 exemples/classe puis full data) entre une "
         "baseline classique et des méthodes modernes, métrique principale F1 macro :")
    bullet(doc, "Baseline : TF-IDF + Régression logistique (ancre de réalité, sans GPU).")
    bullet(doc, "Méthode récente : SetFit (3 backbones), comparée à un fine-tuning CamemBERT, "
                "à un LLM-as-classifier (Ministral 8B via API, zero/few-shot), et à un système "
                "hybride SetFit + fallback Ministral par seuil de confiance.")
    body(doc,
         "Livrables : courbes d'apprentissage, matrice de décision « quel modèle pour quel "
         "profil de club ? », et un dashboard Streamlit simple pour interroger le modèle.")

    doc.save(OUT)
    print(f"Plan prévisionnel généré : {OUT}")


if __name__ == "__main__":
    build()
